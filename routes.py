from flask import Blueprint, render_template, request, jsonify, send_file, Response
import os
from docx import Document
from docx.shared import Inches
import io
import logging
import xarray as xr
import numpy as np
import json
import plotly.graph_objs as go
import plotly.utils
import pandas as pd
import folium
from openpyxl import Workbook
from openpyxl.styles import NamedStyle, Font
import re
from werkzeug.utils import secure_filename

bp = Blueprint('main', __name__)

# Global dataset
current_dataset = {'ds': None, 'filename': None}

UPLOAD_FOLDER = 'uploads'

# Allowed NetCDF file check
def allowed_file(filename):
    """Allow NetCDF (.nc) and GRIB (.grib, .grb) files"""
    allowed_ext = ['nc', 'grib', 'grb']
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed_ext

# Convert numpy types for JSON
def convert_numpy_types(obj):
    if isinstance(obj, np.integer): return int(obj)
    elif isinstance(obj, np.floating): return float(obj)
    elif isinstance(obj, np.ndarray): return obj.tolist()
    elif isinstance(obj, dict): return {k: convert_numpy_types(v) for k, v in obj.items()}
    elif isinstance(obj, list): return [convert_numpy_types(i) for i in obj]
    elif isinstance(obj, tuple): return tuple(convert_numpy_types(i) for i in obj)
    else: return obj

# Extract NetCDF info
def extract_file_info(filepath):
    """Extract basic info from NetCDF or GRIB file"""
    try:
        ext = filepath.rsplit('.', 1)[1].lower()
        if ext == 'nc':
            with xr.open_dataset(filepath, cache=False) as ds:
                current_dataset['ds'] = ds.load()
        elif ext in ['grib', 'grb']:
            with xr.open_dataset(filepath, cache=False, engine='cfgrib') as ds:
                current_dataset['ds'] = ds.load()
        else:
            return {'success': False, 'error': 'Unsupported file type'}


        # Detect a time-like coordinate automatically
        time_var = None
        for coord in current_dataset['ds'].coords:
            c = str(coord).lower()
            if any(x in c for x in ['time', 'date', 'datetime', 'valid_time']):
                time_var = coord
                break

        if time_var:
            time_min = str(current_dataset['ds'][time_var].values.min())
            time_max = str(current_dataset['ds'][time_var].values.max())
            print(f"Detected time coordinate: {time_var}, min: {time_min}, max: {time_max}")
        else:
            print("No time-like coordinate found in dataset")
            time_min = time_max = None



        coords = {}
        for dim in current_dataset['ds'].dims:
            if dim in current_dataset['ds'].coords:
                coords[dim] = {
                    'min': float(current_dataset['ds'].coords[dim].min().values),
                    'max': float(current_dataset['ds'].coords[dim].max().values),
                    'size': int(current_dataset['ds'].sizes[dim])
                }

        variables = {}
        for var in current_dataset['ds'].data_vars:
            attrs = dict(current_dataset['ds'][var].attrs) if hasattr(ds[var], 'attrs') else {}
            variables[var] = {
                'dims': list(current_dataset['ds'][var].dims),
                'shape': [int(x) for x in current_dataset['ds'][var].shape],
                'attrs': convert_numpy_types(attrs)
            }

        global_attrs = dict(current_dataset['ds'].attrs) if hasattr(current_dataset['ds'], 'attrs') else {}

        #ds.close()
        return {
            'coords': coords,
            'variables': variables,
            'global_attrs': convert_numpy_types(global_attrs),
            'success': True
        }
    
    except Exception as e:
        return {'success': False, 'error': str(e)}




def create_coverage_map(ds):
    lat_var = lon_var = None
    for coord in ds.coords:
        c = str(coord).lower()
        if 'lat' in c: lat_var = coord
        elif 'lon' in c: lon_var = coord
    if not lat_var or not lon_var:
        return None

    lats = ds.coords[lat_var].values
    lons = ds.coords[lon_var].values
    center_lat = float(np.mean(lats))
    center_lon = float(np.mean(lons))

    # Map bounds
    south, north = -6, 6
    west, east = 32, 43

    m = folium.Map(
        location=[center_lat, center_lon],
        zoom_start=6,
        max_bounds=True
    )

    # Force map to fit bounds at startup
    m.fit_bounds([[south, west], [north, east]])

    # Enforce strict bounds via JS
    bounds_script = f"""
    <script>
    {m.get_name()}.setMaxBounds([
        [{south}, {west}],
        [{north}, {east}]
    ]);
    </script>
    """
    m.get_root().add_child(folium.Element(bounds_script))

    # Rectangle showing full dataset bounds (optional)
    bounds = [[float(np.min(lats))-0.5, float(np.min(lons))-0.5],
              [float(np.max(lats))+0.5, float(np.max(lons))+0.5]]
    
    south, west = bounds[0]
    north, east = bounds[1]
    
    folium.Rectangle(
        bounds=bounds,
        color= "red",
       fill=False,
        fillOpacity=0.0,
        popup='Data Coverage Area'
    ).add_to(m)

    # Add markers for all points
    for lat in lats:
        for lon in lons:
            folium.CircleMarker(
                location=[float(lat), float(lon)],

                radius=2,
                color="blue",
                fill=False,
                fill_color=None,
                fill_opacity=0,
                weight=0.5,
                tooltip=f"Click: Lat {lat:.4f}, Lon: {lon:.4f}",
            ).add_to(m)

    # Attach one global click binding for all CircleMarkers
    map_name = m.get_name()
    click_script = f"""
    window.addEventListener("load", function() {{
        if (typeof {map_name} !== "undefined") {{
            {map_name}.eachLayer(function(layer) {{
                if (layer instanceof L.CircleMarker) {{
                    layer.on('click', function(e) {{
                        console.log("Marker clicked at", e.latlng); // debug log
                        if (window.parent && typeof window.parent.handleGridClick === 'function') {{
                            window.parent.handleGridClick(e.latlng.lat, e.latlng.lng);
                        }}
                    }});
                }}
            }});
            // Attach to map background (anywhere else)


        }} else {{
            console.error("Map variable {map_name} not defined yet");
        }}
    }});
    """
    # Assuming you already have your lat/lon arrays
    lat_vals = np.array(lats)   # 1D array
    lon_vals = np.array(lons)   # 1D array


    # Estimate average grid spacing
    dlat = np.mean(np.diff(lat_vals))
    dlon = np.mean(np.diff(lon_vals))

    # Use half the diagonal of grid cell as cutoff

    latlon_threshold = np.sqrt((dlat/2)**2 + (dlon/2)**2)
    # Convert to JavaScript arrays
    lat_js = "[" + ",".join(map(str, lat_vals)) + "]"
    lon_js = "[" + ",".join(map(str, lon_vals)) + "]"

    snap_click_script = f"""
    <script>
    
    window.addEventListener("load", function() {{
        if (typeof {m.get_name()} !== "undefined") {{
            var lat_vals = {lat_js};
            var lon_vals = {lon_js};

            function findClosest(arr, val) {{
                return arr.reduce(function(prev, curr) {{
                    return (Math.abs(curr - val) < Math.abs(prev - val) ? curr : prev);
                }});
            }}

            {m.get_name()}.on('click', function(e) {{
            
                clickLat = e.latlng.lat;
                clickLon = e.latlng.lng;
                
                if (clickLat >= {south} && clickLat <= {north} && clickLon >= {west} && clickLon <= {east}) {{
                    console.log("55Clicked inside rectangle at", e.latlng);
                    var nearestLat = findClosest(lat_vals, clickLat);
                var nearestLon = findClosest(lon_vals, clickLon);

                console.log("Clicked:", clickLat.toFixed(4), clickLon.toFixed(4));
                console.log("Snapped to:", nearestLat.toFixed(4), nearestLon.toFixed(4));

                if (window.parent && typeof window.parent.handleGridClick === 'function') {{
                    window.parent.handleGridClick(nearestLat, nearestLon);
                }}
                }} else {{
                    console.log("You clicked outside the rectangle");
                    
                }}

                
            }});
        }} else {{
            console.error("Map variable {m.get_name()} not defined yet");
        }}
    }});
    </script>
    """

    m.get_root().html.add_child(folium.Element(snap_click_script))
    m.get_root().html.add_child(folium.Element(f"<script>{click_script}</script>"))
    


    return m




# Routes
@bp.route('/')
def index():
    #if dataset in memory

    '''ds = current_dataset.get('ds')
    if ds is not None:
        lat_var = lon_var = None
        for coord in ds.coords:
            c = str(coord).lower()
            if 'lat' in c: lat_var = coord
            elif 'lon' in c: lon_var = coord
        lats = ds[lat_var].values.tolist() if lat_var else []
        lons = ds[lon_var].values.tolist() if lon_var else []
    else:
        lats, lons = [], []

    return render_template("index.html", lats=lats, lons=lons)'''

    return render_template('index.html')

@bp.route('/upload', methods=['POST'])
def upload_file():
    try:
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': 'No file provided'})
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'success': False, 'error': 'No file selected'})
        if not allowed_file(file.filename):
            return jsonify({'success': False, 'error': 'Invalid file type. Upload .nc only'})

        filename = secure_filename(file.filename)
        filepath = os.path.join(UPLOAD_FOLDER, filename)
        filepath = os.path.join(UPLOAD_FOLDER, filename)
        if os.path.exists(filepath):
            os.remove(filepath)  # or use os.replace() from a temp file
        file.save(filepath)
                
        info = extract_file_info(filepath)
        if not info['success']:
            return jsonify({'success': False, 'error': info['error']})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})
    try:
        with xr.open_dataset(filepath, cache=False, decode_timedelta=True, chunks="auto") as ds:
            current_dataset['ds'] = ds.load()
            current_dataset['filename'] = filename
            coverage_map = create_coverage_map(current_dataset['ds'])
            map_html = coverage_map._repr_html_() if coverage_map else None


            
            current_dataset['ds'].close()
            return jsonify({'success': True, 'info': info, 'map_html': map_html})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@bp.route('/get_timeseries', methods=['POST'])
def get_timeseries():
    ds = current_dataset.get('ds')
    if ds is None:
        return jsonify({'success': False, 'error': 'No dataset loaded'})

    data = request.get_json()
    lat, lon = data.get('lat'), data.get('lon')
    start = pd.to_datetime(data.get("startDate")) if data.get("startDate") else None
    end   = pd.to_datetime(data.get("endDate")) if data.get("endDate") else None
    print("Json data", data)

    if lat is None or lon is None:
        return jsonify({'success': False, 'error': 'Coordinates not provided'})

    # --- Detect coordinates once ---
    lat_var = lon_var = time_var = None
    for coord in ds.coords:
        c = str(coord).lower()
        if 'lat' in c: 
            lat_var = coord
        elif 'lon' in c: 
            lon_var = coord
        elif 'time' in c or 'date' in c: 
            time_var = coord

    if not lat_var or not lon_var:
        return jsonify({'success': False, 'error': 'Lat/Lon not found in dataset'})

    # --- Slice dataset if time coord + range given ---
    if time_var and start is not None and end is not None:
        # Match tz-awareness
        if pd.api.types.is_datetime64tz_dtype(ds[time_var]):
            if start.tzinfo is None:
                start = start.tz_localize(ds[time_var].dt.tz)
            if end.tzinfo is None:
                end = end.tz_localize(ds[time_var].dt.tz)
        else:
            if start.tzinfo is not None:
                start = start.tz_convert(None)
            if end.tzinfo is not None:
                end = end.tz_convert(None)
  
        ds = ds.sel({time_var: slice(start, end)})
        
        #ds = ds.sel({time_var: (ds[time_var] >= np.datetime64(start)) &
                               #(ds[time_var] <= np.datetime64(end))})
        print(f"Sliced dataset from {start} to end {end} using '{time_var}'")

        if ds[time_var].size == 0:
            return jsonify({'success': False, 'error': 'No data in selected date range'})

        print(f"Detected new time coordinate: {time_var}, min={ds[time_var].values.min()}, max={ds[time_var].values.max()}")
    else:
        print("No time slicing applied; using full dataset")

    # --- Select nearest point ---
    point = ds.sel({lat_var: lat, lon_var: lon}, method='nearest')
    charts = {}

    # --- Build timeseries for each variable ---
    for var_name in ds.data_vars:
        var_data = point[var_name]
        print("Variable:", var_name)

        if var_data.size == 0 or var_data.isnull().all():
            print(f"Skipping {var_name} (empty or all NaN)")
            continue

        # Convert Kelvin → Celsius
        units = var_data.attrs.get('units', '').lower()
        if 'k' in units:
            var_data = var_data - 273.15
            var_data.attrs['units'] = 'C'
            print(f"Converted {var_name} from Kelvin to Celsius")

        if time_var and time_var in var_data.dims:
            times = var_data[time_var].values
            values = var_data.values
            
            if values is not None and values.size > 0:
                print("values are there:", type(values), values.shape)
                print("time values are there also:", type(times), times.shape)

            if len(times) == 0 or len(values) == 0:
                print(f"No data to plot for {var_name}")
                continue

            fig = go.Figure()
            fig.add_trace(go.Scatter(
                x=times,
                y=values,
                mode='lines+markers',
                name=var_name,
                line=dict(color='blue', width=1, dash='solid'),
                marker=dict(symbol='circle', size=2, color='blue')
            ))

            fig.update_layout(
                title=f"Time Series of {var_name}",
                xaxis_title='Time',
                yaxis_title=var_data.attrs.get('units', 'Value'),
                template='plotly',
                legend=dict(x=0.01, y=0.99)
            )

            charts[var_name] = json.loads(plotly.utils.PlotlyJSONEncoder().encode(fig))
    print("Charts dict so far:", charts.keys())
    return jsonify({
        'success': True,
        'charts': charts,
        'coordinates': {
            'lat': float(point[lat_var].values),
            'lon': float(point[lon_var].values)
        }
    })

def sanitize_text(text):
    """Sanitize text to remove non-printable characters and ensure Word compatibility."""
    if text is None or pd.isna(text):
        return "NaN"
    text = str(text)
    # Remove non-printable characters (keep ASCII printable and common Unicode)
    text = re.sub(r'[^\x20-\x7E\u00A0-\uFFFF]', '', text)
    return text

@bp.route("/download_timeseries_csv", methods=["POST"])
def download_timeseries_csv():
    ds = current_dataset.get("ds")
    if ds is None:
        return "No dataset loaded", 400

    data = request.get_json()
    lat, lon = data.get("lat"), data.get("lon")
    start = pd.to_datetime(data.get("startDate")) if data.get("startDate") else None
    end   = pd.to_datetime(data.get("endDate")) + pd.Timedelta(days=1) if data.get("endDate") else None
    print("CHECKING END DATE", end)
    filetype = data.get("filetype", "csv")  # default to CSV if not provided
    keep_constant = data.get("keep_constant", False)  # Option to keep zero-only columns

    # --- Detect coordinates ---
    lat_var = lon_var = time_var = None
    for coord in ds.coords:
        c = str(coord).lower()
        if "lat" in c: lat_var = coord
        elif "lon" in c: lon_var = coord
        elif "time" in c or "date" in c: time_var = coord

    if not lat_var or not lon_var:
        return "Lat/Lon not found in dataset", 400

    # --- Slice time range if provided ---
    if time_var and start is not None and end is not None:
        ds = ds.sel({time_var: slice(start, end)})
        end = end - pd.Timedelta(days=1)
        if ds[time_var].size == 0:
            return "No data in selected date range", 400
    
    # --- Select nearest grid point ---
    point = ds.sel({lat_var: lat, lon_var: lon}, method="nearest")

    # Convert to DataFrame
    df = point.to_dataframe().reset_index()

    # --- Handle units (Kelvin → Celsius as example) ---
    for var_name in point.data_vars:
        units = point[var_name].attrs.get("units", "").lower()
        if "k" in units:
            df[var_name] = df[var_name] - 273.15
            point[var_name].attrs["units"] = "C"

    # --- Filter out columns with all NaN or empty values ---
    df = df.dropna(axis=1, how="all")

    # --- Track and filter constant-value columns ---
    dropped_columns = []
    if not keep_constant:
        # Identify columns with a single unique value
        for col in df.columns:
            if df[col].nunique(dropna=False) == 1:
                # Get the constant value (first non-NaN value, or NaN if all NaN)
                constant_value = df[col].iloc[0] if not df[col].isna().all() else "NaN"
                # Format as Variable_name(constant_value_it_has)
                dropped_columns.append(f"{col}({constant_value})")
        # Drop constant columns
        constant_cols = [col for col in df.columns if df[col].nunique(dropna=False) == 1]
        df = df.drop(columns=constant_cols)


    # --- Filter out columns with all NaN or empty values ---
    df = df.dropna(axis=1, how="all")
    
    # --- Filter out columns with all zeros (if keep_zeros is False) ---
    if not keep_constant:
            # Identify columns with a single unique value (including zeros, ones, or any constant)
            constant_cols = [col for col in df.columns if df[col].nunique(dropna=False) == 1]
            df = df.drop(columns=constant_cols)


    # ==============================
    # 📤 EXPORT SECTION (moved out)
    # ==============================

    
    if filetype == "csv":
        # Generate CSV string from DataFrame
        csv_output = df.to_csv(index=False)
        # Prepend comment row with dropped columns (if any)

        # Build metadata as comment lines (like paragraphs in DOCX)
        header_lines = []
        header_lines.append("# Time Series Data")
        header_lines.append(f"# Grid Point: Lat {lat:.4f}, Lon {lon:.4f}")
        if start and end:
            header_lines.append(f"# Date Range: {start.date()} → {end.date()}")
        if dropped_columns:
            header_lines.append(f"# Dropped constant columns: {' '.join(dropped_columns)}")
        header_lines.append("")  # blank line before CSV table

        # Prepend metadata to CSV
        csv_output = "\n".join(header_lines) + "\n" + csv_output
        return Response(
            csv_output,
            mimetype="text/csv",
            headers={"Content-Disposition": "attachment;filename=timeseries.csv"}
        )
    
    elif filetype == "xlsx":
        # Create workbook
        wb = Workbook()
        ws = wb.active

        # Define styles
        title_style = NamedStyle(name="title_style")
        title_style.font = Font(size=14, bold=True, color="1F497D")  # dark blue

        heading_style = NamedStyle(name="heading_style")
        heading_style.font = Font(size=11, bold=True, color="1F497D")

        # Add metadata with styles
        ws.append(["Time Series Data"])
        ws["A1"].style = title_style

        ws.append([f"Grid Point: Lat {lat:.4f}, Lon {lon:.4f}"])
        ws["A2"].style = heading_style

        if start and end:
            ws.append([f"Date Range: {start.date()} → {end.date()}"])
            ws["A3"].style = heading_style

        if dropped_columns:
            ws.append([f"Dropped constant columns: {' '.join(dropped_columns)}"])
            ws[f"A{ws.max_row}"].style = heading_style

        ws.append([])  # blank row

        # Write table headers + data
        ws.append(df.columns.tolist())
        for row in df.itertuples(index=False, name=None):
            ws.append(row)

        # Set all columns to width 20
        for col in ws.columns:
            col_letter = col[0].column_letter
            ws.column_dimensions[col_letter].width = 20

        buf = io.BytesIO()
        wb.save(buf)
        buf.seek(0)

        return send_file(
            buf,
            mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            as_attachment=True,
            download_name="timeseries.xlsx",
        )
        '''buf = io.StringIO()
        df.to_csv(buf, index=False)
        buf.seek(0)
        return send_file(
            io.BytesIO(buf.getvalue().encode("utf-8")),
            mimetype="text/csv",
            as_attachment=True,
            download_name="timeseries.csv",
        )'''


    elif filetype == "txt":
        # Generate plain text from DataFrame
        buf = io.StringIO()
        df.to_string(buf, index=False)  # Keeps nice table-like formatting
        buf.seek(0)

        # Add metadata (lat/lon, date range, dropped columns) at top
        header = []
        header.append("Time Series Data")
        header.append(f"Grid Point: Lat {lat:.4f}, Lon {lon:.4f}")
        if start and end:
            header.append(f"Date Range: {start.date()} → {end.date()}")
        if dropped_columns:
            header.append(f"Dropped constant columns: {', '.join(dropped_columns)}")
        header.append("")  # Blank line before table

        content = "\n".join(header) + "\n" + buf.getvalue()

        # Return as downloadable .txt file
        return Response(
            content,
            mimetype="text/plain",
            headers={"Content-Disposition": "attachment;filename=timeseries.txt"}
        )
    
    elif filetype == "docx":
        print("before trying:The filetype=docx is called")
        try:
            logging.info("inside try now. elif docx called")
            print("The filetype=docx is called")
            # Create a new Word document
            doc = Document()
            
            # Add metadata as paragraphs
            doc.add_heading("Time Series Data", level=1)
            doc.add_paragraph(f"Grid Point: Lat {lat:.4f}, Lon {lon:.4f}")
            if start and end:
                doc.add_paragraph(f"Date Range: {start.date()} → {end.date()}")
            if dropped_columns:
                doc.add_paragraph(f"Dropped constant columns: {', '.join(dropped_columns)}")
            doc.add_paragraph("")

            # Sanitize DataFrame (column names and values)
            df_clean = df.fillna("NaN").astype(str)
            df_clean.columns = [sanitize_text(col) for col in df_clean.columns]
            for col in df_clean.columns:
                df_clean[col] = df_clean[col].apply(sanitize_text)
            logging.debug(f"DataFrame shape: {df_clean.shape}, columns: {list(df_clean.columns)}")

            # Limit rows to prevent large tables (optional)
            max_rows = 1000
            if len(df_clean) > max_rows:
                logging.warning(f"DataFrame truncated to {max_rows} rows for DOCX")
                df_clean = df_clean.head(max_rows)

            # Create a table
            table = doc.add_table(rows=1 + len(df_clean), cols=len(df_clean.columns))
            table.style = "Light Grid"  # Simpler style for better compatibility

            # Add column headers
            for j, col in enumerate(df_clean.columns):
                cell = table.cell(0, j)
                cell.text = col
                # Set basic formatting to avoid Word issues
                paragraph = cell.paragraphs[0]
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(col)
                run.font.name = "Calibri"
                run.font.size = None  # Default size

            # Add data rows
            for i, row in df_clean.iterrows():
                for j, value in enumerate(row):
                    cell = table.cell(i + 1, j)
                    cell.text = value
                    # Set basic formatting
                    paragraph = cell.paragraphs[0]
                    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(value)
                    run.font.name = "Calibri"
                    run.font.size = None

            # Auto-fit table (instead of fixed widths)
            table.autofit = True

            # Save the document to a BytesIO buffer
            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            docx_content = buf.getvalue()
            logging.debug(f"Generated DOCX file size: {len(docx_content)} bytes")
            buf.close()

            # Save a copy for debugging (optional, remove in production)
            with open("debug_timeseries.docx", "wb") as f:
                f.write(docx_content)
            logging.debug("Saved debug_timeseries.docx for inspection")

            return Response(
                docx_content,
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": "attachment;filename=timeseries.docx"}
            )
        except Exception as e:
            logging.error(f"Error generating DOCX: {str(e)}")
            return "Error generating DOCX file", 500
    else:
        return "Unsupported filetype", 400

